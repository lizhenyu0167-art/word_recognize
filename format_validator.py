#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word文档格式验证工具
功能：验证格式化后的文档是否正确应用了模板样式
"""

import os
import json
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from config import config

class FormatValidator:
    def __init__(self):
        self.template_styles = {}
        self.formatted_styles = {}
        self.validation_report = {}
        
    def analyze_document_styles(self, doc_path):
        """
        分析文档中的样式定义
        """
        try:
            doc = Document(doc_path)
            styles_info = {}
            
            for style in doc.styles:
                if style.type == 1:  # 段落样式
                    style_info = {
                        'style_name': style.name,
                        'style_type': 'paragraph',
                        'font_name': self._get_font_name(style),
                        'font_size': self._get_font_size(style),
                        'bold': self._get_bold(style),
                        'italic': self._get_italic(style)
                    }
                    
                    # 获取字体信息
                    font_info = self.get_font_info(style)
                    style_info.update(font_info)
                    
                    # 获取段落格式信息
                    pf_info = self._get_paragraph_format_info(style)
                    style_info.update(pf_info)
                    
                    styles_info[style.name] = style_info
                    
            return styles_info
            
        except Exception as e:
            print(f"分析文档样式时出错: {e}")
            return {}
    
    def get_font_info(self, style):
        """
        获取字体信息，包括中英文字体分离设置
        """
        font_info = {
            'ascii_font': None,
            'eastasia_font': None,
            'cs_font': None,
            'hansi_font': None
        }
        
        try:
            if hasattr(style, '_element'):
                style_element = style._element
                
                # 查找rPr元素
                rpr = style_element.find(qn('w:rPr'))
                if rpr is not None:
                    # 查找rFonts元素
                    rfonts = rpr.find(qn('w:rFonts'))
                    if rfonts is not None:
                        # 获取各种字体设置
                        ascii_font = rfonts.get(qn('w:ascii'))
                        if ascii_font:
                            font_info['ascii_font'] = ascii_font
                        
                        eastAsia_font = rfonts.get(qn('w:eastAsia'))
                        if eastAsia_font:
                            font_info['eastasia_font'] = eastAsia_font
                        
                        cs_font = rfonts.get(qn('w:cs'))
                        if cs_font:
                            font_info['cs_font'] = cs_font
                        
                        hAnsi_font = rfonts.get(qn('w:hAnsi'))
                        if hAnsi_font:
                            font_info['hansi_font'] = hAnsi_font
                        
        except Exception as e:
            print(f"获取字体信息时出错: {e}")
            
        return font_info

    def _is_font_consistent_with_separation(self, template_style, formatted_style, formatted_font_name):
        """
        检查字体名称是否与字体分离设置一致
        当模板使用继承默认字体时，检查格式化后的字体是否符合字体分离规则
        """
        try:
            # 获取模板的字体分离设置
            template_ascii = template_style.get('ascii_font')
            template_eastasia = template_style.get('eastasia_font')
            
            # 获取格式化后的字体分离设置
            formatted_ascii = formatted_style.get('ascii_font')
            formatted_eastasia = formatted_style.get('eastasia_font')
            
            # 如果模板有明确的字体分离设置
            if template_ascii or template_eastasia:
                # 检查格式化后的字体是否匹配模板的字体分离设置
                # 对于英文内容，应该使用ascii字体
                if template_ascii and formatted_font_name == template_ascii:
                    return True
                # 对于中文内容，应该使用eastasia字体
                if template_eastasia and formatted_font_name == template_eastasia:
                    return True
                # 如果格式化后的字体分离设置与模板一致
                if formatted_ascii == template_ascii and formatted_eastasia == template_eastasia:
                    return True
            
            # 如果没有明确的字体分离设置，使用默认规则
            # 继承默认字体通常对应Times New Roman(英文)或宋体(中文)
            if formatted_font_name in ['Times New Roman', '宋体', 'Arial']:
                return True
                
            return False
            
        except Exception as e:
            print(f"检查字体分离一致性时出错: {e}")
            # 出错时采用宽松匹配
            return formatted_font_name in ['Times New Roman', '宋体', 'Arial']

    def _get_font_name(self, style):
        """获取字体名称"""
        try:
            if hasattr(style, 'font') and style.font.name:
                return style.font.name
            return '继承默认字体'
        except Exception as e:
            print(f"获取字体名称时出错: {e}")
            return None
    
    def _get_font_size(self, style):
        """获取字体大小"""
        try:
            if hasattr(style, 'font') and style.font.size:
                return f"{style.font.size.pt}pt"
            return None
        except Exception as e:
            print(f"获取字体大小时出错: {e}")
            return None
    
    def _get_bold(self, style):
        """获取加粗设置"""
        try:
            if hasattr(style, 'font'):
                return style.font.bold
            return None
        except Exception as e:
            print(f"获取加粗设置时出错: {e}")
            return None
    
    def _get_italic(self, style):
        """获取斜体设置"""
        try:
            if hasattr(style, 'font'):
                return style.font.italic
            return None
        except Exception as e:
            print(f"获取斜体设置时出错: {e}")
            return None
    
    def _get_paragraph_format_info(self, style):
        """获取段落格式信息"""
        info = {}
        
        try:
            pf = style.paragraph_format
            
            # 对齐方式
            if pf.alignment is not None:
                info['alignment'] = f"{pf.alignment.name} ({pf.alignment.value})"
            else:
                info['alignment'] = None
            
            # 行间距
            if pf.line_spacing is not None:
                info['line_spacing'] = str(pf.line_spacing)
            else:
                info['line_spacing'] = None
            
            # 段前距
            if pf.space_before is not None:
                info['space_before'] = f"{pf.space_before.pt}pt"
            else:
                info['space_before'] = None
            
            # 段后距
            if pf.space_after is not None:
                info['space_after'] = f"{pf.space_after.pt}pt"
            else:
                info['space_after'] = None
            
            # 首行缩进
            if pf.first_line_indent is not None:
                info['first_line_indent'] = f"{pf.first_line_indent.pt}pt"
            else:
                info['first_line_indent'] = None
            
            # 左缩进
            if pf.left_indent is not None:
                info['left_indent'] = f"{pf.left_indent.pt}pt"
            else:
                info['left_indent'] = None
                
        except Exception as e:
            print(f"获取段落格式信息时出错: {e}")
            
        return info
    
    def analyze_document_paragraphs(self, doc_path):
        """
        分析文档中段落的实际格式
        """
        try:
            doc = Document(doc_path)
            paragraphs_info = []
            
            for i, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():  # 只分析有内容的段落
                    para_info = {
                        'paragraph_index': i + 1,
                        'text_preview': paragraph.text[:50] + '...' if len(paragraph.text) > 50 else paragraph.text,
                        'style_name': paragraph.style.name,
                        'font_name': self._get_run_font_name(paragraph),
                        'alignment': self._get_paragraph_alignment(paragraph),
                        'line_spacing': self._get_paragraph_line_spacing(paragraph),
                        'first_line_indent': self._get_paragraph_first_line_indent(paragraph)
                    }
                    paragraphs_info.append(para_info)
                    
            return paragraphs_info
            
        except Exception as e:
            print(f"分析文档段落时出错: {e}")
            return []
    
    def _get_run_font_name(self, paragraph):
        """获取段落中运行的字体名称"""
        try:
            for run in paragraph.runs:
                if hasattr(run, 'font') and run.font.name:
                    return run.font.name
            return None
        except Exception as e:
            print(f"获取运行字体名称时出错: {e}")
            return None
    
    def _get_paragraph_alignment(self, paragraph):
        """获取段落对齐方式"""
        try:
            if hasattr(paragraph, 'alignment') and paragraph.alignment is not None:
                return f"{paragraph.alignment.name} ({paragraph.alignment.value})"
            return None
        except Exception as e:
            print(f"获取段落对齐方式时出错: {e}")
            return None
    
    def _get_paragraph_line_spacing(self, paragraph):
        """获取段落行间距"""
        try:
            if hasattr(paragraph, 'paragraph_format') and paragraph.paragraph_format.line_spacing is not None:
                return str(paragraph.paragraph_format.line_spacing)
            return None
        except Exception as e:
            print(f"获取段落行间距时出错: {e}")
            return None
    
    def _get_paragraph_first_line_indent(self, paragraph):
        """获取段落首行缩进"""
        try:
            if hasattr(paragraph, 'paragraph_format') and paragraph.paragraph_format.first_line_indent is not None:
                return f"{paragraph.paragraph_format.first_line_indent.pt}pt"
            return None
        except Exception as e:
            print(f"获取段落首行缩进时出错: {e}")
            return None
    
    def compare_styles(self, template_styles, formatted_styles):
        """
        比较模板样式和格式化后样式的差异
        """
        comparison_result = {}
        
        for style_name in template_styles:
            if style_name in formatted_styles:
                template_style = template_styles[style_name]
                formatted_style = formatted_styles[style_name]
                
                matches = []
                differences = []
                
                # 比较所有属性
                all_properties = set(template_style.keys()) | set(formatted_style.keys())
                
                for prop in all_properties:
                    template_value = template_style.get(prop)
                    formatted_value = formatted_style.get(prop)
                    
                    # 如果模板中该属性为None，则跳过比较（表示未定义）
                    if template_value is None:
                        continue
                    
                    # 特殊处理字体名称：考虑字体分离机制
                    if prop == 'font_name' and template_value == '继承默认字体':
                        # 当模板使用继承默认字体时，检查字体分离设置
                        if self._is_font_consistent_with_separation(template_style, formatted_style, formatted_value):
                            matches.append(prop)
                        else:
                            differences.append({
                                'property': prop,
                                'template_value': template_value,
                                'formatted_value': formatted_value,
                                'note': '字体分离机制导致的差异'
                            })
                    elif template_value == formatted_value:
                        matches.append(prop)
                    else:
                        differences.append({
                            'property': prop,
                            'template_value': template_value,
                            'formatted_value': formatted_value
                        })
                
                comparison_result[style_name] = {
                    'status': 'matched' if len(differences) == 0 else 'different',
                    'matches': matches,
                    'differences': differences
                }
            else:
                comparison_result[style_name] = {
                    'status': 'missing',
                    'matches': [],
                    'differences': [{'property': 'entire_style', 'template_value': 'exists', 'formatted_value': 'missing'}]
                }
        
        return comparison_result
    
    def generate_validation_report(self, template_doc, formatted_doc, output_file=None):
        """
        生成完整的验证报告
        """
        if output_file is None:
            output_file = config.VALIDATION_REPORT
        
        print("开始格式验证...")
        
        # 确保输出目录存在
        config.ensure_output_dir()
        
        # 分析模板样式
        print("分析格式模板样式...")
        template_styles = self.analyze_document_styles(template_doc)
        
        # 分析格式化后文档样式
        print("分析格式化后文档样式...")
        formatted_styles = self.analyze_document_styles(formatted_doc)
        
        # 分析格式化后文档段落
        formatted_paragraphs = self.analyze_document_paragraphs(formatted_doc)
        
        # 比较样式差异
        print("比较样式差异...")
        style_comparison = self.compare_styles(template_styles, formatted_styles)
        
        # 生成报告
        report = {
            'template_document': template_doc,
            'formatted_document': formatted_doc,
            'template_styles_count': len(template_styles),
            'formatted_styles_count': len(formatted_styles),
            'paragraphs_count': len(formatted_paragraphs),
            'style_comparison': style_comparison,
            'template_styles': template_styles,
            'formatted_styles': formatted_styles,
            'formatted_paragraphs': formatted_paragraphs
        }
        
        # 保存详细报告
        with open(output_file, 'w', encoding='utf-8') as f:
            json.dump(report, f, ensure_ascii=False, indent=2)
        
        # 生成控制台摘要
        self._print_validation_summary(report)
        
        return report
    
    def _print_validation_summary(self, report):
        """
        打印验证摘要到控制台
        """
        print("\n=== 格式验证报告 ===")
        print(f"格式模板: {report['template_document']}")
        print(f"格式化后文档: {report['formatted_document']}")
        print(f"模板样式数量: {report['template_styles_count']}")
        print(f"格式化后样式数量: {report['formatted_styles_count']}")
        print(f"文档段落数量: {report['paragraphs_count']}")
        
        print("\n=== 样式比较结果 ===")
        
        matched_count = 0
        different_count = 0
        missing_count = 0
        
        for style_name, comparison in report['style_comparison'].items():
            status = comparison['status']
            
            if status == 'matched':
                print(f"✅ {style_name}: 完全匹配")
                matched_count += 1
            elif status == 'different':
                diff_count = len(comparison['differences'])
                print(f"⚠️  {style_name}: 存在差异 ({diff_count}个属性不同)")
                
                # 显示前几个差异
                for diff in comparison['differences'][:3]:  # 只显示前3个差异
                    template_val = diff['template_value'] if diff['template_value'] is not None else 'None'
                    formatted_val = diff['formatted_value'] if diff['formatted_value'] is not None else 'None'
                    print(f"   - {diff['property']}: 模板='{template_val}' vs 格式化='{formatted_val}'")
                
                different_count += 1
            elif status == 'missing':
                print(f"❌ {style_name}: 缺少样式")
                missing_count += 1
        
        print("\n=== 总结 ===")
        print(f"完全匹配: {matched_count}个样式")
        print(f"存在差异: {different_count}个样式")
        print(f"缺少样式: {missing_count}个样式")
        
        total_styles = matched_count + different_count + missing_count
        if total_styles > 0:
            match_rate = (matched_count / total_styles) * 100
            print(f"格式匹配率: {match_rate:.1f}%")
            
            if match_rate >= 80:
                print("✅ 格式转换效果优秀！")
            elif match_rate >= 60:
                print("⚠️  格式转换基本成功，但仍有部分样式需要调整。")
            else:
                print("❌ 格式转换存在较多问题，需要检查转换逻辑。")

def main():
    # 文件路径
    template_doc = config.TEMPLATE_FILE
    
    # 使用配置中的格式化文档路径
    formatted_doc = config.get_formatted_doc_path()
    
    if not formatted_doc:
        print(f"错误: 找不到格式化后文档（模式: {config.FORMATTED_DOC_PREFIX}*.docx）")
        print("请先运行 format_applier.py 生成格式化文档")
        return
    
    # 检查文件是否存在
    if not os.path.exists(template_doc):
        print(f"错误: 找不到格式模板文件 {template_doc}")
        return
    
    if not os.path.exists(formatted_doc):
        print(f"错误: 找不到格式化后文档 {formatted_doc}")
        print("请先运行 format_applier.py 生成格式化文档")
        return
    
    # 创建验证器
    validator = FormatValidator()
    
    # 生成验证报告
    report = validator.generate_validation_report(template_doc, formatted_doc)
    
    print(f"\n详细报告已保存到: {config.VALIDATION_REPORT}")

if __name__ == "__main__":
    main()