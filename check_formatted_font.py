#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
检查格式化后文档的字体设置
"""

from docx import Document
from docx.oxml.ns import qn

def check_formatted_document():
    """
    检查格式化后文档的字体设置
    """
    doc_path = "output/格式化后的测试文档.docx"
    
    try:
        doc = Document(doc_path)
        print(f"检查文档: {doc_path}")
        print(f"段落数量: {len(doc.paragraphs)}")
        
        # 检查Normal样式的设置
        print("\n=== Normal样式检查 ===")
        for style in doc.styles:
            if style.name == 'Normal':
                print(f"样式名称: {style.name}")
                if hasattr(style, 'font'):
                    print(f"font.name: {style.font.name}")
                    print(f"font.size: {style.font.size}")
                
                # 检查XML级别的字体分离设置
                if hasattr(style, '_element'):
                    style_element = style._element
                    rpr = style_element.find(qn('w:rPr'))
                    if rpr is not None:
                        rfonts = rpr.find(qn('w:rFonts'))
                        if rfonts is not None:
                            ascii_font = rfonts.get(qn('w:ascii'))
            
                            hAnsi_font = rfonts.get(qn('w:hAnsi'))
                            eastAsia_font = rfonts.get(qn('w:eastAsia'))
                            cs_font = rfonts.get(qn('w:cs'))
                            print(f"XML字体分离:")
                            print(f"  ascii: {ascii_font}")
                            print(f"  hAnsi: {hAnsi_font}")
                            print(f"  eastAsia: {eastAsia_font}")
                            print(f"  cs: {cs_font}")
                        else:
                            print("未找到rFonts元素")
                    else:
                        print("未找到rPr元素")
                break
        
        # 检查前几个段落的字体设置
        print("\n=== 段落字体检查 ===")
        for i, paragraph in enumerate(doc.paragraphs[:5]):
            print(f"\n段落 {i+1}: '{paragraph.text[:50]}...'")
            print(f"样式: {paragraph.style.name if paragraph.style else '无'}")
            
            # 检查段落级别的字体设置
            if hasattr(paragraph, '_element'):
                para_element = paragraph._element
                ppr = para_element.find(qn('w:pPr'))
                if ppr is not None:
                    para_rpr = ppr.find(qn('w:rPr'))
                    if para_rpr is not None:
                        para_rfonts = para_rpr.find(qn('w:rFonts'))
                        if para_rfonts is not None:
                            print("  段落级别字体设置存在（应该已被清除）")
                        else:
                            print("  段落级别无字体设置 ✓")
                    else:
                        print("  段落级别无字体设置 ✓")
                else:
                    print("  段落级别无字体设置 ✓")
            
            # 检查run级别的字体设置
            for j, run in enumerate(paragraph.runs):
                if hasattr(run, '_element'):
                    run_element = run._element
                    rpr = run_element.find(qn('w:rPr'))
                    if rpr is not None:
                        rfonts = rpr.find(qn('w:rFonts'))
                        if rfonts is not None:
                            print(f"  Run {j+1} 字体设置存在（应该已被清除）")
                        else:
                            print(f"  Run {j+1} 无字体设置 ✓")
                    else:
                        print(f"  Run {j+1} 无字体设置 ✓")
                        
    except Exception as e:
        print(f"检查文档时出错: {e}")

if __name__ == "__main__":
    check_formatted_document()