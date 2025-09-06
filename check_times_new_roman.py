#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
检查Word文档中Times New Roman字体的设置情况
"""

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import parse_xml
import os

def check_times_new_roman_settings():
    """检查Word文档中Times New Roman字体的设置"""
    doc_path = "格式模板.docx"
    
    if not os.path.exists(doc_path):
        print(f"文档不存在: {doc_path}")
        return
    
    doc = Document(doc_path)
    print(f"检查文档: {doc_path}")
    print("="*50)
    
    # 1. 检查文档默认字体设置
    print("\n=== 文档默认字体设置 ===")
    try:
        doc_defaults = doc.part.element.xpath('//w:docDefaults')[0]
        rpr_default = doc_defaults.xpath('.//w:rPrDefault/w:rPr')[0]
        
        # 检查默认字体
        rfonts = rpr_default.xpath('.//w:rFonts')[0] if rpr_default.xpath('.//w:rFonts') else None
        if rfonts is not None:
            ascii_font = rfonts.get(qn('w:ascii'))
            hansi_font = rfonts.get(qn('w:hAnsi'))
            eastasia_font = rfonts.get(qn('w:eastAsia'))
            cs_font = rfonts.get(qn('w:cs'))
            
            print(f"  默认ascii字体: {ascii_font or '未设置'}")
            print(f"  默认hAnsi字体: {hansi_font or '未设置'}")
            print(f"  默认eastAsia字体: {eastasia_font or '未设置'}")
            print(f"  默认cs字体: {cs_font or '未设置'}")
        else:
            print("  无默认字体设置")
    except Exception as e:
        print(f"  获取文档默认设置失败: {e}")
    
    # 2. 检查所有样式中的Times New Roman设置
    print("\n=== 样式中的Times New Roman设置 ===")
    times_new_roman_styles = []
    
    for style in doc.styles:
        try:
            if hasattr(style, '_element') and style._element is not None:
                # 检查样式的字体设置
                rpr_elements = style._element.xpath('.//w:rPr')
                for rpr in rpr_elements:
                    rfonts_elements = rpr.xpath('.//w:rFonts')
                    for rfonts in rfonts_elements:
                        ascii_font = rfonts.get(qn('w:ascii'))
                        hansi_font = rfonts.get(qn('w:hAnsi'))
                        eastasia_font = rfonts.get(qn('w:eastAsia'))
                        cs_font = rfonts.get(qn('w:cs'))
                        
                        if any(font == "Times New Roman" for font in [ascii_font, hansi_font, eastasia_font, cs_font]):
                            times_new_roman_styles.append({
                                'style_name': style.name,
                                'ascii': ascii_font,
                                'hAnsi': hansi_font,
                                'eastAsia': eastasia_font,
                                'cs': cs_font
                            })
        except Exception as e:
            print(f"  检查样式 {style.name} 失败: {e}")
    
    if times_new_roman_styles:
        print(f"  找到 {len(times_new_roman_styles)} 个使用Times New Roman的样式:")
        for style_info in times_new_roman_styles:
            print(f"    样式: {style_info['style_name']}")
            print(f"      ascii: {style_info['ascii'] or '未设置'}")
            print(f"      hAnsi: {style_info['hAnsi'] or '未设置'}")
            print(f"      eastAsia: {style_info['eastAsia'] or '未设置'}")
            print(f"      cs: {style_info['cs'] or '未设置'}")
            print()
    else:
        print("  未找到使用Times New Roman的样式")
    
    # 3. 检查段落中的直接字体设置
    print("\n=== 段落中的Times New Roman设置 ===")
    times_new_roman_paragraphs = []
    
    for i, paragraph in enumerate(doc.paragraphs):
        try:
            # 检查段落级别的字体设置
            if paragraph._element.xpath('.//w:rFonts'):
                for rfonts in paragraph._element.xpath('.//w:rFonts'):
                    ascii_font = rfonts.get(qn('w:ascii'))
                    hansi_font = rfonts.get(qn('w:hAnsi'))
                    eastasia_font = rfonts.get(qn('w:eastAsia'))
                    cs_font = rfonts.get(qn('w:cs'))
                    
                    if any(font == "Times New Roman" for font in [ascii_font, hansi_font, eastasia_font, cs_font]):
                        times_new_roman_paragraphs.append({
                            'paragraph_index': i,
                            'text': paragraph.text[:50] + "..." if len(paragraph.text) > 50 else paragraph.text,
                            'ascii': ascii_font,
                            'hAnsi': hansi_font,
                            'eastAsia': eastasia_font,
                            'cs': cs_font
                        })
        except Exception as e:
            print(f"  检查段落 {i} 失败: {e}")
    
    if times_new_roman_paragraphs:
        print(f"  找到 {len(times_new_roman_paragraphs)} 个使用Times New Roman的段落:")
        for para_info in times_new_roman_paragraphs:
            print(f"    段落 {para_info['paragraph_index']}: {para_info['text']}")
            print(f"      ascii: {para_info['ascii'] or '未设置'}")
            print(f"      hAnsi: {para_info['hAnsi'] or '未设置'}")
            print(f"      eastAsia: {para_info['eastAsia'] or '未设置'}")
            print(f"      cs: {para_info['cs'] or '未设置'}")
            print()
    else:
        print("  未找到使用Times New Roman的段落")
    
    # 4. 检查运行级别的字体设置
    print("\n=== 运行级别的Times New Roman设置 ===")
    times_new_roman_runs = []
    
    for para_i, paragraph in enumerate(doc.paragraphs):
        for run_i, run in enumerate(paragraph.runs):
            try:
                if run._element.xpath('.//w:rFonts'):
                    for rfonts in run._element.xpath('.//w:rFonts'):
                        ascii_font = rfonts.get(qn('w:ascii'))
                        hansi_font = rfonts.get(qn('w:hAnsi'))
                        eastasia_font = rfonts.get(qn('w:eastAsia'))
                        cs_font = rfonts.get(qn('w:cs'))
                        
                        if any(font == "Times New Roman" for font in [ascii_font, hansi_font, eastasia_font, cs_font]):
                            times_new_roman_runs.append({
                                'paragraph_index': para_i,
                                'run_index': run_i,
                                'text': run.text[:30] + "..." if len(run.text) > 30 else run.text,
                                'ascii': ascii_font,
                                'hAnsi': hansi_font,
                                'eastAsia': eastasia_font,
                                'cs': cs_font
                            })
            except Exception as e:
                print(f"  检查运行 {para_i}-{run_i} 失败: {e}")
    
    if times_new_roman_runs:
        print(f"  找到 {len(times_new_roman_runs)} 个使用Times New Roman的运行:")
        for run_info in times_new_roman_runs:
            print(f"    段落 {run_info['paragraph_index']}, 运行 {run_info['run_index']}: {run_info['text']}")
            print(f"      ascii: {run_info['ascii'] or '未设置'}")
            print(f"      hAnsi: {run_info['hAnsi'] or '未设置'}")
            print(f"      eastAsia: {run_info['eastAsia'] or '未设置'}")
            print(f"      cs: {run_info['cs'] or '未设置'}")
            print()
    else:
        print("  未找到使用Times New Roman的运行")
    
    print("\n检查完成！")

if __name__ == "__main__":
    check_times_new_roman_settings()