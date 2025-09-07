#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document

# 加载格式模板文档
print("检查格式模板页眉内容:")
doc = Document('格式模板.docx')
for i, section in enumerate(doc.sections):
    print(f"第{i+1}节页眉:")
    for p in section.header.paragraphs:
        print(f"  - '{p.text}'")

# 加载测试文档
print("\n检查测试文档页眉内容:")
test_doc = Document('测试文档.docx')
for i, section in enumerate(test_doc.sections):
    print(f"第{i+1}节页眉:")
    for p in section.header.paragraphs:
        print(f"  - '{p.text}'")

# 检查奇偶页设置
print("\n检查奇偶页设置:")
for i, section in enumerate(doc.sections):
    print(f"格式模板第{i+1}节: 奇偶页不同={section.odd_and_even_pages_header_footer}")

for i, section in enumerate(test_doc.sections):
    print(f"测试文档第{i+1}节: 奇偶页不同={section.odd_and_even_pages_header_footer}")