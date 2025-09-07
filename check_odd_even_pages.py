#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
import os
from config import config

# 加载格式化后的文档
formatted_doc_path = config.get_fixed_formatted_doc_path()
print(f"检查文档: {formatted_doc_path}")

if not os.path.exists(formatted_doc_path):
    print(f"错误：文件不存在 {formatted_doc_path}")
    exit(1)

doc = Document(formatted_doc_path)

# 检查奇偶页设置 - 使用XML级别检查
print("\n检查奇偶页设置:")
from docx.oxml.ns import qn

for i, section in enumerate(doc.sections):
    # 在XML级别检查奇偶页设置
    section_element = section._sectPr
    even_and_odd_headers = section_element.find(qn('w:evenAndOddHeaders'))
    odd_even_enabled = even_and_odd_headers is not None
    print(f"第{i+1}节: 奇偶页不同={odd_even_enabled}")
    
    # 检查奇数页页眉内容
    print(f"  奇数页页眉内容:")
    for p in section.header.paragraphs:
        if p.text.strip():
            print(f"    - '{p.text}'")
    
    # 检查偶数页页眉内容
    print(f"  偶数页页眉内容:")
    for p in section.even_page_header.paragraphs:
        if p.text.strip():
            print(f"    - '{p.text}'")

print("\n验证完成")