#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
检查HTML Preformatted样式
"""

from docx import Document
import os

def check_html_preformatted_style():
    """检查HTML Preformatted样式是否存在于文档中"""
    
    # 检查原始文档
    original_doc_path = "测试文档.docx"
    if os.path.exists(original_doc_path):
        print(f"=== 检查原始文档: {original_doc_path} ===")
        doc = Document(original_doc_path)
        
        print("\n文档中的所有样式:")
        for style in doc.styles:
            print(f"  - {style.name} (类型: {style.type})")
        
        # 检查是否有HTML Preformatted样式
        html_preformatted_found = False
        for style in doc.styles:
            if style.name == "HTML Preformatted":
                html_preformatted_found = True
                print(f"\n✅ 找到HTML Preformatted样式: {style}")
                break
        
        if not html_preformatted_found:
            print("\n❌ 原始文档中没有HTML Preformatted样式")
    
    # 检查格式化后的文档
    formatted_doc_path = "output/格式化后的测试文档.docx"
    if os.path.exists(formatted_doc_path):
        print(f"\n=== 检查格式化后文档: {formatted_doc_path} ===")
        doc = Document(formatted_doc_path)
        
        print("\n文档中的所有样式:")
        for style in doc.styles:
            print(f"  - {style.name} (类型: {style.type})")
        
        # 检查是否有HTML Preformatted样式
        html_preformatted_found = False
        for style in doc.styles:
            if style.name == "HTML Preformatted":
                html_preformatted_found = True
                print(f"\n✅ 找到HTML Preformatted样式: {style}")
                break
        
        if not html_preformatted_found:
            print("\n❌ 格式化后文档中没有HTML Preformatted样式")
    
    # 检查模板文档
    template_doc_path = "格式模板.docx"
    if os.path.exists(template_doc_path):
        print(f"\n=== 检查模板文档: {template_doc_path} ===")
        doc = Document(template_doc_path)
        
        print("\n文档中的所有样式:")
        for style in doc.styles:
            print(f"  - {style.name} (类型: {style.type})")
        
        # 检查是否有HTML Preformatted样式
        html_preformatted_found = False
        for style in doc.styles:
            if style.name == "HTML Preformatted":
                html_preformatted_found = True
                print(f"\n✅ 找到HTML Preformatted样式: {style}")
                print(f"样式类型: {style.type}")
                if hasattr(style, 'font'):
                    print(f"字体名称: {style.font.name}")
                    print(f"字体大小: {style.font.size}")
                break
        
        if not html_preformatted_found:
            print("\n❌ 模板文档中没有HTML Preformatted样式")

if __name__ == "__main__":
    check_html_preformatted_style()