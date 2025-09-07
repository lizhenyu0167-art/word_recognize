#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
动态格式应用器
功能：基于动态提取的格式信息，将格式模板的格式应用到测试文档
完全移除硬编码映射规则，实现真正的动态格式应用
包括页眉页脚格式的应用
"""

import os
import json
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX, WD_UNDERLINE
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from datetime import datetime
from config import config

class DynamicFormatApplier:
    def __init__(self, format_info_path=None):
        self.format_info_path = format_info_path or config.DYNAMIC_FORMAT_INFO
        self.format_info = None
        self.alignment_map = {
            '左对齐': WD_ALIGN_PARAGRAPH.LEFT,
            '居中': WD_ALIGN_PARAGRAPH.CENTER,
            '右对齐': WD_ALIGN_PARAGRAPH.RIGHT,
            '两端对齐': WD_ALIGN_PARAGRAPH.JUSTIFY,
            '分散对齐': WD_ALIGN_PARAGRAPH.DISTRIBUTE
        }
    
    def load_format_info(self, format_file=None):
        """
        加载动态提取的格式信息
        """
        if format_file is None:
            format_file = self.format_info_path
        
        try:
            if os.path.exists(format_file):
                with open(format_file, 'r', encoding='utf-8') as f:
                    self.format_info = json.load(f)
                print(f"已加载格式信息: {format_file}")
                print(f"模板文件: {self.format_info.get('template_file', '未知')}")
                print(f"提取时间: {self.format_info.get('extraction_time', '未知')}")
                print(f"样式数量: {len(self.format_info.get('styles', {}))}")
                return True
            else:
                print(f"格式信息文件不存在: {format_file}")
                return False
                
        except Exception as e:
            print(f"加载格式信息时出错: {e}")
            return False
    
    def apply_formats_to_document(self, input_path=None, output_path=None, use_clean_document=True):
        """
        将动态格式信息应用到测试文档
        """
        if not self.format_info:
            print("错误：未加载格式信息，请先调用load_format_info()")
            return False
        
        if input_path is None:
            # 选择使用清理后的文档还是原始文档
            if use_clean_document:
                clean_doc_path = os.path.join(config.OUTPUT_DIR, "测试文档_清理后.docx")
                if os.path.exists(clean_doc_path):
                    input_path = clean_doc_path
                    print(f"\n使用清理后的测试文档: {input_path}")
                else:
                    input_path = config.TEST_DOCUMENT
                    print(f"\n清理后的文档不存在，使用原始测试文档: {input_path}")
            else:
                input_path = config.TEST_DOCUMENT
                print(f"\n使用原始测试文档: {input_path}")
        
        if output_path is None:
            output_path = config.get_fixed_formatted_doc_path()
        
        print(f"\n正在应用格式到文档: {input_path}")
        
        try:
            doc = Document(input_path)
            
            # 1. 应用文档默认设置
            self._apply_document_defaults(doc)
            
            # 2. 应用样式格式
            print("\n=== 应用样式格式 ===")
            
            # 首先加载模板文档以便复制缺失的样式
            template_doc = None
            template_path = self.format_info.get('template_file')
            if template_path and os.path.exists(template_path):
                try:
                    template_doc = Document(template_path)
                    print(f"已加载模板文档用于样式复制: {template_path}")
                except Exception as e:
                    print(f"警告：无法加载模板文档 {template_path}: {e}")
            
            for style_name, style_info in self.format_info['styles'].items():
                # 检查样式是否存在，如果不存在则尝试从模板复制
                if not self._ensure_style_exists(doc, style_name, template_doc):
                    continue
                    
                if self._apply_style_format(doc, style_name, style_info):
                    print(f"已应用样式: {style_name}")
                    
                    # 显示字体分离信息
                    if 'font_separation' in style_info:
                        sep = style_info['font_separation']
                        ascii_font = sep.get('ascii', '未设置')
                        eastAsia_font = sep.get('eastAsia', '未设置')
                        print(f"  字体分离: 英文={ascii_font}, 中文={eastAsia_font}")
            
            # 3. 应用页眉页脚格式
            self._apply_header_footer_formats(doc, input_path)
            
            # 4. 清除段落级别的字体设置，让段落继承样式字体
            self._clear_paragraph_fonts(doc)
            
            # 5. 保存格式化后的文档
            doc.save(output_path)
            print(f"\n格式化完成！文档已保存为: {output_path}")
            return True
            
        except Exception as e:
            print(f"应用格式时出错: {e}")
            return False
    
    def _apply_document_defaults(self, doc):
        """
        应用文档默认设置
        """
        try:
            defaults = self.format_info.get('document_defaults', {})
            default_font = defaults.get('default_font')
            default_font_size = defaults.get('default_font_size')
            
            if default_font or default_font_size:
                print(f"应用文档默认设置: 字体={default_font}, 字号={default_font_size}")
                
                # 注意：不在这里设置Normal样式的font.name，
                # 因为这会覆盖后续的字体分离设置
                # 字体分离设置会在_apply_style_format中处理
            
            # 为所有节设置奇偶页不同的页眉页脚 - 在XML级别设置
            from docx.oxml.ns import qn
            for section in doc.sections:
                # 在XML级别设置奇偶页不同
                section_element = section._sectPr
                even_and_odd_headers = section_element.find(qn('w:evenAndOddHeaders'))
                if even_and_odd_headers is None:
                    even_and_odd_headers = section_element.makeelement(qn('w:evenAndOddHeaders'), {})
                    section_element.append(even_and_odd_headers)
            print("已设置文档默认使用奇偶页不同的页眉页脚（XML级别）")
                        
        except Exception as e:
            print(f"应用文档默认设置时出错: {e}")
    
    def _ensure_style_exists(self, doc, style_name, template_doc=None):
        """
        确保样式存在，如果不存在则尝试从模板复制
        """
        # 检查样式是否已存在
        for style in doc.styles:
            if style.name == style_name:
                return True
        
        # 样式不存在，尝试从模板复制
        if template_doc is not None:
            template_style = None
            for style in template_doc.styles:
                if style.name == style_name:
                    template_style = style
                    break
            
            if template_style is not None:
                try:
                    print(f"  从模板复制样式: {style_name}")
                    # 根据样式类型添加新样式
                    if template_style.type == 1:  # PARAGRAPH
                        new_style = doc.styles.add_style(style_name, 1)
                    elif template_style.type == 2:  # CHARACTER
                        new_style = doc.styles.add_style(style_name, 2)
                    else:
                        print(f"  警告：不支持的样式类型 {template_style.type} for {style_name}")
                        return False
                    
                    # 复制基本属性
                    if hasattr(template_style, 'base_style') and template_style.base_style:
                        try:
                            new_style.base_style = template_style.base_style
                        except:
                            pass  # 如果基础样式不存在，忽略错误
                    
                    return True
                    
                except Exception as e:
                    print(f"  警告：无法复制样式 {style_name}: {e}")
                    return False
        
        print(f"  警告：未找到样式 {style_name}")
        return False
    
    def _apply_style_format(self, doc, style_name, style_info):
        """
        应用单个样式的格式
        """
        try:
            # 查找对应的样式
            target_style = None
            for style in doc.styles:
                if style.name == style_name:
                    target_style = style
                    break
            
            if not target_style:
                print(f"  错误：样式 {style_name} 不存在")
                return False
            
            # 应用字体格式
            self._apply_font_format(target_style, style_info)
            
            # 应用段落格式
            self._apply_paragraph_format(target_style, style_info)
            
            # 应用字体分离设置（XML级别）
            if 'font_separation' in style_info:
                self._apply_font_separation(target_style, style_info['font_separation'])
            
            return True
            
        except Exception as e:
            print(f"应用样式 {style_name} 格式时出错: {e}")
            return False
    
    def _apply_font_format(self, style, style_info):
        """
        应用字体格式
        """
        try:
            if hasattr(style, 'font'):
                font = style.font
                
                # 基本字体名称 - 只有在明确设置且不是"继承默认字体"时才设置
                if 'font_name' in style_info and style_info['font_name'] != '继承默认字体':
                    font.name = style_info['font_name']
                elif 'font_name' in style_info and style_info['font_name'] == '继承默认字体':
                    # 清除font.name，让它继承默认设置
                    # 但保留字体分离设置，因为它们可能包含重要的语言特定字体信息
                    font.name = None
                    # 注意：不清除XML级别的字体分离设置，因为它们是独立的配置
                    # 字体分离设置将通过_apply_font_separation方法单独处理
                
                # 字号
                if 'font_size' in style_info:
                    size_pt = float(style_info['font_size'].replace('pt', ''))
                    font.size = Pt(size_pt)
                
                # 粗体
                if 'bold' in style_info:
                    font.bold = style_info['bold']
                
                # 斜体
                if 'italic' in style_info:
                    font.italic = style_info['italic']
                
                # 下划线
                if 'underline' in style_info:
                    # 这里可以根据需要处理下划线设置
                    pass
                
                # 字体颜色
                if 'color' in style_info:
                    try:
                        from docx.shared import RGBColor
                        color_str = style_info['color']
                        # 处理十六进制颜色格式，如 'EE0000'
                        if len(color_str) == 6:
                            r = int(color_str[0:2], 16)
                            g = int(color_str[2:4], 16)
                            b = int(color_str[4:6], 16)
                            font.color.rgb = RGBColor(r, g, b)
                            print(f"  已应用字体颜色: RGB({r}, {g}, {b})")
                        else:
                            print(f"警告：无法解析颜色格式: {color_str}")
                    except Exception as color_error:
                        print(f"应用字体颜色时出错: {color_error}")
                
        except Exception as e:
            print(f"应用字体格式时出错: {e}")
    
    def _apply_paragraph_format(self, style, style_info):
        """
        应用段落格式
        """
        try:
            if hasattr(style, 'paragraph_format'):
                pf = style.paragraph_format
                
                # 对齐方式
                if 'alignment' in style_info:
                    alignment_value = self.alignment_map.get(style_info['alignment'])
                    if alignment_value is not None:
                        pf.alignment = alignment_value
                else:
                    # 清除对齐方式，使用默认值
                    pf.alignment = None
                
                # 行间距 - 完整覆盖逻辑
                if 'line_spacing' in style_info:
                    pf.line_spacing = float(style_info['line_spacing'])
                else:
                    # 清除行距设置，使用默认值
                    pf.line_spacing = None
                
                # 段前距
                if 'space_before' in style_info:
                    space_pt = float(style_info['space_before'].replace('pt', ''))
                    pf.space_before = Pt(space_pt)
                else:
                    # 清除段前距，使用默认值
                    pf.space_before = None
                
                # 段后距
                if 'space_after' in style_info:
                    space_pt = float(style_info['space_after'].replace('pt', ''))
                    pf.space_after = Pt(space_pt)
                else:
                    # 清除段后距，使用默认值
                    pf.space_after = None
                
                # 首行缩进
                if 'first_line_indent' in style_info:
                    indent_pt = float(style_info['first_line_indent'].replace('pt', ''))
                    pf.first_line_indent = Pt(indent_pt)
                else:
                    # 清除首行缩进，使用默认值
                    pf.first_line_indent = None
                
                # 左缩进
                if 'left_indent' in style_info:
                    indent_pt = float(style_info['left_indent'].replace('pt', ''))
                    pf.left_indent = Pt(indent_pt)
                else:
                    # 清除左缩进，使用默认值
                    pf.left_indent = None
                
                # 右缩进
                if 'right_indent' in style_info:
                    indent_pt = float(style_info['right_indent'].replace('pt', ''))
                    pf.right_indent = Pt(indent_pt)
                else:
                    # 清除右缩进，使用默认值
                    pf.right_indent = None
                    
        except Exception as e:
            print(f"应用段落格式时出错: {e}")
    
    def _apply_font_separation(self, style, font_separation):
        """
        应用字体分离设置（XML级别）
        """
        try:
            if hasattr(style, '_element'):
                style_element = style._element
                
                # 查找或创建rPr元素
                rpr = style_element.find(qn('w:rPr'))
                if rpr is None:
                    rpr = style_element.makeelement(qn('w:rPr'))
                    style_element.insert(0, rpr)
                
                # 查找或创建rFonts元素
                rfonts = rpr.find(qn('w:rFonts'))
                if rfonts is None:
                    rfonts = rpr.makeelement(qn('w:rFonts'))
                    rpr.insert(0, rfonts)
                
                # 智能应用字体分离设置
                # 只设置明确指定的字体，未指定的字体清除以继承默认设置
                
                # 设置指定的字体
                if 'ascii' in font_separation and font_separation['ascii'] != '未设置':
                    rfonts.set(qn('w:ascii'), font_separation['ascii'])
                else:
                    # 清除ascii字体设置，让它继承默认
                    if qn('w:ascii') in rfonts.attrib:
                        rfonts.attrib.pop(qn('w:ascii'), None)
                
                if 'hAnsi' in font_separation and font_separation['hAnsi'] != '未设置':
                    rfonts.set(qn('w:hAnsi'), font_separation['hAnsi'])
                else:
                    # 清除hAnsi字体设置，让它继承默认
                    if qn('w:hAnsi') in rfonts.attrib:
                        rfonts.attrib.pop(qn('w:hAnsi'), None)
                
                if 'eastAsia' in font_separation and font_separation['eastAsia'] != '未设置':
                    rfonts.set(qn('w:eastAsia'), font_separation['eastAsia'])
                else:
                    # 清除eastAsia字体设置，让它继承默认
                    if qn('w:eastAsia') in rfonts.attrib:
                        rfonts.attrib.pop(qn('w:eastAsia'), None)
                
                if 'cs' in font_separation and font_separation['cs'] != '未设置':
                    rfonts.set(qn('w:cs'), font_separation['cs'])
                else:
                    # 清除cs字体设置，让它继承默认
                    if qn('w:cs') in rfonts.attrib:
                        rfonts.attrib.pop(qn('w:cs'), None)
                    
        except Exception as e:
            print(f"应用字体分离设置时出错: {e}")
    
    def _clear_paragraph_fonts(self, doc):
        """
        清除段落级别的字体设置，让段落继承样式字体
        """
        try:
            print("\n=== 清除段落级别字体设置 ===")
            cleared_run_count = 0
            cleared_para_count = 0
            
            for paragraph in doc.paragraphs:
                # 清除段落级别的字体设置
                if hasattr(paragraph, '_element'):
                    para_element = paragraph._element
                    
                    # 查找段落的pPr元素
                    ppr = para_element.find(qn('w:pPr'))
                    if ppr is not None:
                        # 查找并移除段落级别的rPr元素
                        para_rpr = ppr.find(qn('w:rPr'))
                        if para_rpr is not None:
                            # 移除段落级别的rFonts
                            para_rfonts = para_rpr.find(qn('w:rFonts'))
                            if para_rfonts is not None:
                                para_rpr.remove(para_rfonts)
                                cleared_para_count += 1
                            
                            # 如果段落rPr为空，也移除它
                            if len(para_rpr) == 0:
                                ppr.remove(para_rpr)
                
                # 清除run级别的字体设置
                for run in paragraph.runs:
                    if hasattr(run, '_element'):
                        run_element = run._element
                        
                        # 查找rPr元素
                        rpr = run_element.find(qn('w:rPr'))
                        if rpr is not None:
                            # 查找并移除rFonts元素
                            rfonts = rpr.find(qn('w:rFonts'))
                            if rfonts is not None:
                                rpr.remove(rfonts)
                                cleared_run_count += 1
                            
                            # 如果rPr为空，也移除它
                            if len(rpr) == 0:
                                run_element.remove(rpr)
            
            print(f"已清除 {cleared_run_count} 个run的字体设置")
            print(f"已清除 {cleared_para_count} 个段落的字体设置")
            
        except Exception as e:
            print(f"清除段落字体设置时出错: {e}")
    
    def _apply_header_footer_formats(self, doc, test_doc_path):
        """
        应用页眉页脚格式
        根据需求：
        - 奇数页的页眉内容为格式模板的页眉内容
        - 偶数页的页眉内容为测试文档的标题一的内容
        - 页脚设置页码，确保奇偶页都有页码
        """
        try:
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            
            print("\n=== 应用页眉页脚格式 ===")
            
            # 加载测试文档以获取标题一内容
            test_doc = Document(test_doc_path)
            
            # 查找标题一内容
            title_one_content = ""
            for para in test_doc.paragraphs:
                if para.style.name == "Heading 1" or para.style.name == "标题 1":
                    title_one_content = para.text
                    break
            
            if not title_one_content:
                print("警告：未找到标题一内容，将使用文档标题作为替代")
                title_one_content = test_doc.core_properties.title or "文档标题"
            
            print(f"找到标题一内容: {title_one_content}")
            
            # 加载格式模板以获取页眉内容
            template_doc = Document(self.format_info.get('template_file') or config.TEMPLATE_FILE)
            
            # 获取格式模板的页眉内容
            odd_header_content = ""
            if len(template_doc.sections) > 0:
                section = template_doc.sections[0]
                for para in section.header.paragraphs:
                    if para.text.strip():
                        odd_header_content = para.text
                        break
            
            if not odd_header_content:
                print("警告：未找到格式模板页眉内容，将使用默认页眉")
                odd_header_content = "社会保障评论"
            
            print(f"找到格式模板页眉内容: {odd_header_content}")
            
            # 设置奇偶页不同的页眉
            for i, section in enumerate(doc.sections):
                # 设置页眉页脚选项
                section.different_first_page_header_footer = False  # 不使用首页不同的页眉
                
                # 设置页眉顶端距离和页脚底端距离
                section_id = f"section_{i+1}"
                section_settings = self.format_info.get('section_settings', {}).get(section_id, {})
                
                # 设置页眉顶端距离
                if 'header_distance' in section_settings:
                    header_distance_str = section_settings['header_distance']
                    if header_distance_str.endswith('pt'):
                        header_distance_pt = float(header_distance_str.replace('pt', ''))
                        section.header_distance = Pt(header_distance_pt)
                        print(f"设置第{i+1}节页眉顶端距离: {header_distance_pt}pt")
                else:
                    section.header_distance = Pt(15)  # 默认设置页眉与正文的距离
                
                # 设置页脚底端距离
                if 'footer_distance' in section_settings:
                    footer_distance_str = section_settings['footer_distance']
                    if footer_distance_str.endswith('pt'):
                        footer_distance_pt = float(footer_distance_str.replace('pt', ''))
                        section.footer_distance = Pt(footer_distance_pt)
                        print(f"设置第{i+1}节页脚底端距离: {footer_distance_pt}pt")
                else:
                    section.footer_distance = Pt(15)  # 默认设置页脚与正文的距离
                
                # 启用奇偶页不同的页眉页脚 - 在XML级别设置
                section_element = section._sectPr
                even_and_odd_headers = section_element.find(qn('w:evenAndOddHeaders'))
                if even_and_odd_headers is None:
                    even_and_odd_headers = section_element.makeelement(qn('w:evenAndOddHeaders'), {})
                    section_element.append(even_and_odd_headers)
                print(f"已启用第{i+1}节的奇偶页不同页眉页脚")
                
                # 清除现有页眉内容
                for para in section.header.paragraphs:
                    for run in para.runs:
                        run.text = ""
                
                # 设置奇数页页眉
                if len(section.header.paragraphs) == 0:
                    para = section.header.add_paragraph()
                else:
                    para = section.header.paragraphs[0]
                    # 清除段落中的所有内容
                    para.clear()
                
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # 确保段落没有缩进
                para.paragraph_format.left_indent = Inches(0)
                para.paragraph_format.right_indent = Inches(0)
                para.paragraph_format.first_line_indent = Inches(0)
                
                # 保留原始格式，包括TAB字符
                # 只去除前导和尾随空格，保留中间的TAB
                content = odd_header_content.strip()
                run = para.add_run(content)
                run.font.size = Pt(10.5)
                
                # 设置偶数页页眉
                if len(section.even_page_header.paragraphs) == 0:
                    para = section.even_page_header.add_paragraph()
                else:
                    para = section.even_page_header.paragraphs[0]
                    # 清除段落中的所有内容
                    para.clear()
                
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # 确保段落没有缩进
                para.paragraph_format.left_indent = Inches(0)
                para.paragraph_format.right_indent = Inches(0)
                para.paragraph_format.first_line_indent = Inches(0)
                
                # 保留原始格式，包括TAB字符
                # 只去除前导和尾随空格，保留中间的TAB
                content = title_one_content.strip()
                run = para.add_run(content)
                run.font.size = Pt(10.5)
                
                # 设置奇数页页脚（页码）
                # 先清除现有页脚内容
                for i in range(len(section.footer.paragraphs)):
                    section.footer.paragraphs[i].clear()
                
                # 删除所有段落并创建新段落
                if len(section.footer.paragraphs) > 0:
                    para = section.footer.paragraphs[0]
                    para.clear()
                else:
                    para = section.footer.add_paragraph()
                
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run("- ")
                run.font.size = Pt(10.5)
                
                # 添加页码域代码（奇数页）
                fldChar = OxmlElement('w:fldChar')
                fldChar.set(qn('w:fldCharType'), 'begin')
                
                instrText = OxmlElement('w:instrText')
                instrText.set(qn('xml:space'), 'preserve')
                instrText.text = " PAGE "
                
                fldChar2 = OxmlElement('w:fldChar')
                fldChar2.set(qn('w:fldCharType'), 'end')
                
                r_element = run._r
                r_element.append(fldChar)
                r_element.append(instrText)
                r_element.append(fldChar2)
                
                run = para.add_run(" -")
                run.font.size = Pt(10.5)
                
                # 设置偶数页页脚（页码）
                # 先清除现有页脚内容
                for i in range(len(section.even_page_footer.paragraphs)):
                    section.even_page_footer.paragraphs[i].clear()
                
                # 删除所有段落并创建新段落
                if len(section.even_page_footer.paragraphs) > 0:
                    para = section.even_page_footer.paragraphs[0]
                    para.clear()
                else:
                    para = section.even_page_footer.add_paragraph()
                
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run("- ")
                run.font.size = Pt(10.5)
                
                # 添加页码域代码
                fldChar = OxmlElement('w:fldChar')
                fldChar.set(qn('w:fldCharType'), 'begin')
                
                instrText = OxmlElement('w:instrText')
                instrText.set(qn('xml:space'), 'preserve')
                instrText.text = " PAGE "
                
                fldChar2 = OxmlElement('w:fldChar')
                fldChar2.set(qn('w:fldCharType'), 'end')
                
                r_element = run._r
                r_element.append(fldChar)
                r_element.append(instrText)
                r_element.append(fldChar2)
                
                run = para.add_run(" -")
                run.font.size = Pt(10.5)
                
                print(f"已设置第{i+1}节的奇偶页页眉和页脚页码")
            
            print("页眉页脚格式应用完成")
            
        except Exception as e:
            print(f"应用页眉页脚格式时出错: {e}")
    
    def get_style_summary(self):
        """
        获取样式摘要信息
        """
        if not self.format_info:
            return "未加载格式信息"
        
        summary = []
        summary.append(f"格式模板: {self.format_info.get('template_file', '未知')}")
        summary.append(f"提取时间: {self.format_info.get('extraction_time', '未知')}")
        summary.append(f"文档默认字体: {self.format_info.get('document_defaults', {}).get('default_font', '未设置')}")
        summary.append(f"样式数量: {len(self.format_info.get('styles', {}))}")
        summary.append(f"页眉数量: {len(self.format_info.get('headers', {}))}")
        summary.append(f"页脚数量: {len(self.format_info.get('footers', {}))}")
        
        # 添加节设置摘要
        section_settings = self.format_info.get('section_settings', {})
        if section_settings:
            summary.append("\n节设置:")
            for section_id, settings in section_settings.items():
                section_summary = [f"  {section_id}:"]
                if 'header_distance' in settings:
                    section_summary.append(f"    页眉顶端距离: {settings['header_distance']}")
                if 'footer_distance' in settings:
                    section_summary.append(f"    页脚底端距离: {settings['footer_distance']}")
                summary.append("\n".join(section_summary))
        
        # 显示主要标题样式的字体分离信息
        styles = self.format_info.get('styles', {})
        for heading_name in ['Heading 1', 'Heading 2', 'Heading 3', 'Heading 4']:
            if heading_name in styles:
                style_info = styles[heading_name]
                if 'font_separation' in style_info:
                    sep = style_info['font_separation']
                    ascii_font = sep.get('ascii', '未设置')
                    eastAsia_font = sep.get('eastAsia', '未设置')
                    summary.append(f"{heading_name}: 英文={ascii_font}, 中文={eastAsia_font}")
        
        return "\n".join(summary)

def main():
    """
    主函数：应用动态格式到测试文档
    """
    # 检查格式信息文件是否存在
    if not os.path.exists(config.DYNAMIC_FORMAT_INFO):
        print(f"错误：找不到格式信息文件 {config.DYNAMIC_FORMAT_INFO}")
        print("请先运行 dynamic_format_extractor.py 提取格式信息")
        return
    
    # 验证必需文件
    missing_files = config.validate_required_files()
    if missing_files:
        print(f"错误：找不到以下必需文件: {', '.join(missing_files)}")
        return
    
    applier = DynamicFormatApplier()
    
    # 1. 加载格式信息
    if not applier.load_format_info():
        print("无法加载格式信息，请先运行 dynamic_format_extractor.py")
        return
    
    print("\n=== 格式信息摘要 ===")
    print(applier.get_style_summary())
    
    # 2. 应用格式到测试文档
    success = applier.apply_formats_to_document()
    if success:
        print("\n格式应用完成！")
    else:
        print("\n格式应用失败！")

if __name__ == "__main__":
    main()