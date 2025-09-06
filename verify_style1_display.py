#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
验证style_1样式在Word中的实际显示效果
"""

from docx import Document
from docx.shared import Pt

def verify_style1_display():
    """
    创建一个测试文档来验证style_1样式的字体显示
    """
    print("=== 验证style_1样式的字体显示 ===")
    
    # 创建新文档
    doc = Document()
    
    # 加载格式模板以获取样式
    template_doc = Document('格式模板.docx')
    
    # 查找style_1样式
    style_1 = None
    for style in template_doc.styles:
        if style.name == 'style_1':
            style_1 = style
            break
    
    if not style_1:
        print("未找到style_1样式")
        return
    
    # 将style_1样式复制到新文档
    try:
        # 复制样式定义
        new_style = doc.styles.add_style('style_1', style_1.type)
        
        # 复制字体设置
        if style_1.font.name:
            new_style.font.name = style_1.font.name
        if style_1.font.size:
            new_style.font.size = style_1.font.size
        if style_1.font.bold is not None:
            new_style.font.bold = style_1.font.bold
        
        # 复制段落格式
        if hasattr(style_1, 'paragraph_format'):
            if style_1.paragraph_format.first_line_indent:
                new_style.paragraph_format.first_line_indent = style_1.paragraph_format.first_line_indent
            if style_1.paragraph_format.left_indent:
                new_style.paragraph_format.left_indent = style_1.paragraph_format.left_indent
        
        print("style_1样式已复制到新文档")
        
    except Exception as e:
        print(f"复制样式时出错: {e}")
        return
    
    # 添加测试内容
    doc.add_paragraph("这是使用style_1样式的测试段落 - This is a test paragraph using style_1 style", style='style_1')
    doc.add_paragraph("English text: Hello World! 中文文本：你好世界！", style='style_1')
    doc.add_paragraph("Mixed content: ABC 中文 123 English 测试", style='style_1')
    
    # 保存测试文档
    output_path = 'output/style_1_test.docx'
    doc.save(output_path)
    print(f"测试文档已保存到: {output_path}")
    
    # 分析原始样式的字体设置
    print("\n=== style_1样式的字体分析 ===")
    print(f"基本字体名称: {style_1.font.name}")
    print(f"字体大小: {style_1.font.size}")
    print(f"是否加粗: {style_1.font.bold}")
    
    # 检查实际的字体分离设置
    from docx.oxml.ns import qn
    if hasattr(style_1, '_element'):
        style_element = style_1._element
        rpr = style_element.find(qn('w:rPr'))
        if rpr is not None:
            rfonts = rpr.find(qn('w:rFonts'))
            if rfonts is not None:
                print("\n实际字体分离设置:")
                for attr in ['ascii', 'hAnsi', 'eastAsia', 'cs']:
                    font_val = rfonts.get(qn(f'w:{attr}'))
                    if font_val:
                        print(f"  {attr}: {font_val}")
                    else:
                        print(f"  {attr}: 未设置")
                
                print("\n主题字体设置:")
                for attr in ['asciiTheme', 'hAnsiTheme', 'eastAsiaTheme', 'cstheme']:
                    theme_val = rfonts.get(qn(f'w:{attr}'))
                    if theme_val:
                        print(f"  {attr}: {theme_val}")
                    else:
                        print(f"  {attr}: 未设置")
            else:
                print("无字体分离设置")
        else:
            print("无运行属性设置")
    
    print("\n请打开生成的测试文档查看style_1样式的实际显示效果")
    print("如果英文和中文显示不同的字体，说明Word内部有字体分离设置")
    print("如果显示相同字体，说明确实没有设置英文字体分离")

if __name__ == "__main__":
    verify_style1_display()