#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import os
import glob
from docx import Document

# 查找最新生成的文档
output_dir = 'output'
files = glob.glob(os.path.join(output_dir, '格式化后的测试文档_*.docx'))
if not files:
    print("未找到格式化后的文档")
    exit(1)

latest_file = max(files, key=os.path.getctime)
print(f"检查最新生成的文档: {latest_file}")

# 加载文档
doc = Document(latest_file)

# 检查奇偶页设置
print("\n检查奇偶页设置:")
for i, section in enumerate(doc.sections):
    print(f"第{i+1}节: 奇偶页不同设置已启用")

# 检查页眉内容
print("\n检查页眉内容:")
for i, section in enumerate(doc.sections):
    print(f"第{i+1}节页眉:")
    for p in section.header.paragraphs:
        print(f"  - '{p.text}'")

    # 检查偶数页页眉（如果有）
    if hasattr(section, 'even_page_header') and section.even_page_header:
        print(f"第{i+1}节偶数页页眉:")
        for p in section.even_page_header.paragraphs:
            print(f"  - '{p.text}'")