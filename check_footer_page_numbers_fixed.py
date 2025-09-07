import os
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def check_footer_page_numbers(doc_path):
    print(f"检查文档: {doc_path}")
    doc = Document(doc_path)
    
    for i, section in enumerate(doc.sections):
        print(f"\n第{i+1}节:")
        # 检查奇偶页不同设置
        print(f"首页不同设置: {section.different_first_page_header_footer}")
        
        # 检查奇偶页不同设置
        sectPr = section._sectPr
        evenAndOddHeaders = sectPr.xpath('./w:evenAndOddHeaders')
        print(f"奇偶页不同设置: {len(evenAndOddHeaders) > 0}")
        
        # 检查奇数页页脚
        print("\n奇数页页脚内容:")
        for j, para in enumerate(section.footer.paragraphs):
            print(f"  段落 {j+1}: {para.text}")
            
            # 检查是否有页码域代码
            has_field = False
            for run in para.runs:
                # 检查XML元素中是否有域代码
                r_xml = run._element.xml
                if "fldChar" in r_xml and "instrText" in r_xml and "PAGE" in r_xml:
                    has_field = True
                    print(f"  - 包含页码域代码")
                    # 打印run的文本内容
                    print(f"  - Run文本: '{run.text}'")
            
            # 检查段落中的所有run
            print(f"  - 段落中的run数量: {len(para.runs)}")
            for k, run in enumerate(para.runs):
                print(f"    Run {k+1}: '{run.text}'")
        
        # 检查偶数页页脚
        print("\n偶数页页脚内容:")
        for j, para in enumerate(section.even_page_footer.paragraphs):
            print(f"  段落 {j+1}: {para.text}")
            
            # 检查是否有页码域代码
            has_field = False
            for run in para.runs:
                # 检查XML元素中是否有域代码
                r_xml = run._element.xml
                if "fldChar" in r_xml and "instrText" in r_xml and "PAGE" in r_xml:
                    has_field = True
                    print(f"  - 包含页码域代码")
                    # 打印run的文本内容
                    print(f"  - Run文本: '{run.text}'")
            
            # 检查段落中的所有run
            print(f"  - 段落中的run数量: {len(para.runs)}")
            for k, run in enumerate(para.runs):
                print(f"    Run {k+1}: '{run.text}'")

if __name__ == "__main__":
    output_dir = "output"
    doc_path = os.path.join(output_dir, "格式化后的测试文档.docx")
    
    if os.path.exists(doc_path):
        check_footer_page_numbers(doc_path)
    else:
        print(f"文件不存在: {doc_path}")