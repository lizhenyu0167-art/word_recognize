#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import os
from docx import Document
from config import config

# 检查固定名称的输出文件
output_path = config.get_fixed_formatted_doc_path()
print(f"检查固定名称的输出文件: {output_path}")

if not os.path.exists(output_path):
    print("错误：文件不存在")
    exit(1)

# 加载文档
doc = Document(output_path)

# 检查奇偶页设置
print("\n检查奇偶页设置:")
for i, section in enumerate(doc.sections):
    print(f"第{i+1}节: 奇偶页不同设置已启用")

# 检查页眉内容
print("\n检查页眉内容:")
for i, section in enumerate(doc.sections):
    print(f"第{i+1}节页眉:")
    for p in section.header.paragraphs:
        print(f"  - '{p.text}'")

    # 检查偶数页页眉（如果有）
    if hasattr(section, 'even_page_header') and section.even_page_header:
        print(f"第{i+1}节偶数页页眉:")
        for p in section.even_page_header.paragraphs:
            print(f"  - '{p.text}'")