import os
import sys
from docx import Document
import re

def check_header_spacing(doc_path):
    print(f"检查文档: {doc_path}")
    doc = Document(doc_path)
    
    for i, section in enumerate(doc.sections):
        print(f"\n第{i+1}节:")
        
        # 检查奇数页页眉
        print("\n奇数页页眉内容:")
        for j, para in enumerate(section.header.paragraphs):
            # 显示原始文本，包括空格和制表符
            text_with_spaces = para.text
            # 将制表符可视化显示
            text_with_visible_tabs = text_with_spaces.replace('\t', '[TAB]')
            print(f"  段落 {j+1}: '{text_with_visible_tabs}'")
            # 检查段落中的所有run
            print(f"  - 段落中的run数量: {len(para.runs)}")
            for k, run in enumerate(para.runs):
                run_text_with_visible_tabs = run.text.replace('\t', '[TAB]')
                print(f"    Run {k+1}: '{run_text_with_visible_tabs}'")
        
        # 检查偶数页页眉
        print("\n偶数页页眉内容:")
        for j, para in enumerate(section.even_page_header.paragraphs):
            # 显示原始文本，包括空格和制表符
            text_with_spaces = para.text
            # 将制表符可视化显示
            text_with_visible_tabs = text_with_spaces.replace('\t', '[TAB]')
            print(f"  段落 {j+1}: '{text_with_visible_tabs}'")
            # 检查段落中的所有run
            print(f"  - 段落中的run数量: {len(para.runs)}")
            for k, run in enumerate(para.runs):
                run_text_with_visible_tabs = run.text.replace('\t', '[TAB]')
                print(f"    Run {k+1}: '{run_text_with_visible_tabs}'")

if __name__ == "__main__":
    # 如果提供了命令行参数，使用它作为文档路径
    if len(sys.argv) > 1:
        doc_path = sys.argv[1]
        if os.path.exists(doc_path):
            check_header_spacing(doc_path)
        else:
            print(f"文件不存在: {doc_path}")
    else:
        # 默认检查格式化后的文档
        output_dir = "output"
        doc_path = os.path.join(output_dir, "格式化后的测试文档.docx")
        
        if os.path.exists(doc_path):
            check_header_spacing(doc_path)
        else:
            print(f"文件不存在: {doc_path}")